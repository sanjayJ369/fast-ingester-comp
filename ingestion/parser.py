"""
ingestion/parser.py — Parallel document parser.

Supports: PDF, DOCX, TXT, HTML, CSV, Excel (XLSX/XLS)
All parsing is async — 200 docs fire concurrently via asyncio.gather().
"""

import asyncio
import aiofiles
import io
import csv
from pathlib import Path
from dataclasses import dataclass, field
from typing import Optional

import pdfplumber
from docx import Document as DocxDocument
from bs4 import BeautifulSoup


# ── Data model ────────────────────────────────────────────────────────────────

@dataclass
class ParsedDocument:
    doc_id:   str
    filename: str
    pages:    list[str]          # one entry per page / logical section
    metadata: dict = field(default_factory=dict)

    @property
    def full_text(self) -> str:
        return "\n\n".join(self.pages)


# ── Router ────────────────────────────────────────────────────────────────────

async def parse_document(filepath: str | Path) -> Optional[ParsedDocument]:
    """
    Top-level async router. Detects file type and dispatches to the
    correct parser. Returns None on failure (never raises).
    """
    path = Path(filepath)
    suffix = path.suffix.lower()

    try:
        if suffix == ".pdf":
            return await _parse_pdf(path)
        elif suffix in (".docx", ".doc"):
            return await _parse_docx(path)
        elif suffix in (".html", ".htm"):
            return await _parse_html(path)
        elif suffix == ".csv":
            return await _parse_csv(path)
        elif suffix in (".xlsx", ".xls"):
            return await _parse_excel(path)
        else:
            # .txt and everything else
            return await _parse_txt(path)
    except Exception as e:
        print(f"[PARSER] ⚠️  Failed to parse {path.name}: {e}")
        return None


# ── PDF ───────────────────────────────────────────────────────────────────────

async def _parse_pdf(path: Path) -> ParsedDocument:
    """
    Run pdfplumber in a thread pool to avoid blocking the event loop.
    Extracts text per page + tables as TSV rows.
    """
    def _extract():
        pages = []
        with pdfplumber.open(str(path)) as pdf:
            for page in pdf.pages:
                text = page.extract_text() or ""

                # also pull tables and append as plain text rows
                tables = page.extract_tables()
                for table in tables:
                    for row in table:
                        row_text = "\t".join(str(c) for c in row if c)
                        text += "\n" + row_text

                if text.strip():
                    pages.append(text.strip())
        return pages

    loop = asyncio.get_running_loop()
    pages = await loop.run_in_executor(None, _extract)

    return ParsedDocument(
        doc_id=path.stem,
        filename=path.name,
        pages=pages,
        metadata={"type": "pdf", "num_pages": len(pages)},
    )


# ── DOCX ──────────────────────────────────────────────────────────────────────

async def _parse_docx(path: Path) -> ParsedDocument:
    def _extract():
        doc = DocxDocument(str(path))
        sections: list[str] = []
        current: list[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            # treat Heading styles as section boundaries
            if para.style.name.startswith("Heading"):
                if current:
                    sections.append("\n".join(current))
                    current = []
                current.append(f"## {text}")
            else:
                current.append(text)

        if current:
            sections.append("\n".join(current))

        # also extract tables
        for table in doc.tables:
            rows = []
            for row in table.rows:
                rows.append("\t".join(cell.text.strip() for cell in row.cells))
            sections.append("\n".join(rows))

        return sections

    loop = asyncio.get_running_loop()
    pages = await loop.run_in_executor(None, _extract)

    return ParsedDocument(
        doc_id=path.stem,
        filename=path.name,
        pages=pages,
        metadata={"type": "docx"},
    )


# ── HTML ──────────────────────────────────────────────────────────────────────

async def _parse_html(path: Path) -> ParsedDocument:
    async with aiofiles.open(str(path), encoding="utf-8", errors="ignore") as f:
        raw = await f.read()

    soup = BeautifulSoup(raw, "html.parser")
    # remove scripts / styles
    for tag in soup(["script", "style", "nav", "footer"]):
        tag.decompose()

    text = soup.get_text(separator="\n")
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    cleaned = "\n".join(lines)

    return ParsedDocument(
        doc_id=path.stem,
        filename=path.name,
        pages=[cleaned],
        metadata={"type": "html"},
    )


# ── CSV ───────────────────────────────────────────────────────────────────────

async def _parse_csv(path: Path) -> ParsedDocument:
    async with aiofiles.open(str(path), encoding="utf-8", errors="ignore") as f:
        raw = await f.read()

    reader = csv.reader(io.StringIO(raw))
    rows = ["\t".join(row) for row in reader]
    # chunk CSV into 50-row pages to avoid massive single strings
    pages = ["\n".join(rows[i:i+50]) for i in range(0, len(rows), 50)]

    return ParsedDocument(
        doc_id=path.stem,
        filename=path.name,
        pages=pages,
        metadata={"type": "csv", "num_rows": len(rows)},
    )


# ── Excel ─────────────────────────────────────────────────────────────────────

async def _parse_excel(path: Path) -> ParsedDocument:
    """Parse .xlsx/.xls files. Each sheet becomes one or more pages (50 rows each)."""
    def _extract():
        suffix = path.suffix.lower()
        pages = []

        if suffix == ".xlsx":
            import openpyxl
            wb = openpyxl.load_workbook(str(path), data_only=True, read_only=True)
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                rows = []
                for row in ws.iter_rows(values_only=True):
                    row_text = "\t".join(str(c) if c is not None else "" for c in row)
                    if row_text.strip():
                        rows.append(row_text)
                if rows:
                    for i in range(0, len(rows), 50):
                        page = f"[Sheet: {sheet_name}]\n" + "\n".join(rows[i:i+50])
                        pages.append(page)
            wb.close()
        else:  # .xls
            import xlrd
            wb = xlrd.open_workbook(str(path))
            for sheet_name in wb.sheet_names():
                ws = wb.sheet_by_name(sheet_name)
                rows = []
                for rx in range(ws.nrows):
                    row_text = "\t".join(str(c) for c in ws.row_values(rx))
                    if row_text.strip():
                        rows.append(row_text)
                if rows:
                    for i in range(0, len(rows), 50):
                        page = f"[Sheet: {sheet_name}]\n" + "\n".join(rows[i:i+50])
                        pages.append(page)

        return pages

    loop = asyncio.get_running_loop()
    pages = await loop.run_in_executor(None, _extract)

    return ParsedDocument(
        doc_id=path.stem,
        filename=path.name,
        pages=pages,
        metadata={"type": "excel", "num_sheets": len(pages)},
    )


# ── TXT ───────────────────────────────────────────────────────────────────────

async def _parse_txt(path: Path) -> ParsedDocument:
    async with aiofiles.open(str(path), encoding="utf-8", errors="ignore") as f:
        text = await f.read()

    # split on double newlines as paragraph boundaries → natural pages
    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
    # group into ~500 char pages
    pages, current, count = [], [], 0
    for para in paragraphs:
        current.append(para)
        count += len(para)
        if count >= 500:
            pages.append("\n\n".join(current))
            current, count = [], 0
    if current:
        pages.append("\n\n".join(current))

    return ParsedDocument(
        doc_id=path.stem,
        filename=path.name,
        pages=pages or [text[:2000]],
        metadata={"type": "txt"},
    )


# ── Batch entry point ─────────────────────────────────────────────────────────

async def parse_all(filepaths: list[str | Path]) -> list[ParsedDocument]:
    """
    Parse all documents concurrently.
    Usage:
        docs = await parse_all(list_of_paths)
    """
    tasks = [parse_document(fp) for fp in filepaths]
    results = await asyncio.gather(*tasks)
    docs = [r for r in results if r is not None]
    print(f"[PARSER] ✅ Parsed {len(docs)}/{len(filepaths)} documents")
    return docs
